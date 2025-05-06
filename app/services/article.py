import yaml
from langchain_core.messages import SystemMessage, HumanMessage
from app.llm.llm_services import gpt
from docx import Document
import uuid
import os
def generate_article(prompt: str) -> str:
    with open("app/prompts/prompt.yml", "r", encoding="utf-8") as file:
        config = yaml.safe_load(file)
        system_prompt = config["ARTICLE_CREATION_PROMPT"]["prompt"]
    
    system_prompt= system_prompt.format(user_prompt=prompt)
    messages = [
        SystemMessage(content=system_prompt),
        HumanMessage(content=prompt)
    ]
    response = gpt.invoke(messages)
    return response.content



def generate_docx(title: str, content: str, author: str) -> str:
    article_id = str(uuid.uuid4())
    print(f"Generating docx for article with ID {article_id}:")
    print(f"Title: {title}")
    print(f"Content: {content[:100]}...")  
    print(f"Author: {author}")
    directory = "static/articles"
    if not os.path.exists(directory):
        os.makedirs(directory)  
        print(f"Created directory: {directory}")
    doc = Document()
    print("Document object created.")
    
    doc.add_heading(title, 0)
    print(f"Added heading: {title}")
    
    doc.add_paragraph(content)
    print(f"Added content: {content[:100]}...")
    
    doc.add_paragraph(f"Author: {author}")
    print(f"Added author: {author}")
    
    file_path = f"static/articles/article_{article_id}.docx"
    print(f"File path set to: {file_path}")
    try:
        doc.save(file_path)
        print(f"Document saved successfully at {file_path}")
    except Exception as e:
        print(f"Error saving document: {e}")
        raise e 
    return file_path
